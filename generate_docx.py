#!/usr/bin/env python3
"""Generate METHODS.docx from structured content."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(16)
    elif level == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)


def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def add_equation(text):
    """Add equation as indented italic paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table


# ===================== CONTENT =====================

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HFD Heat Field Simulator — 方法与实现文档')
run.bold = True
run.font.size = Pt(18)

add_para('本文档详细描述 HFD（Heat Field Deformation）热场模拟器的物理模型、数值方法、数据处理算法及可视化实现，供论文撰写参考。',
         italic=True, size=11)

# ========== 1 物理模型 ==========
doc.add_heading('1  物理模型', level=1)

doc.add_heading('1.1  控制方程', level=2)
add_para('模拟器求解的是二维各向异性对流-扩散方程（2D anisotropic advection-diffusion equation），描述含液流边材中的热量传输：')
add_equation('∂T/∂t = D_tg · ∂²T/∂x² + D_ax · ∂²T/∂y² − v_th · ∂T/∂y + q/(ρc·Δx²) · δ(x_H, y_H)')

add_table(
    ['符号', '含义', '单位'],
    [
        ['T', '温度', '°C'],
        ['x', '切向坐标（tangential）', 'm'],
        ['y', '轴向坐标（axial，液流方向）', 'm'],
        ['D_tg', '切向热扩散系数', 'm²/s'],
        ['D_ax', '轴向热扩散系数', 'm²/s'],
        ['v_th', '热速度（thermal velocity）', 'm/s'],
        ['q', '线热源功率密度（经端部效应修正）', 'W/m'],
        ['ρc', '边材容积热容', 'J/(m³·K)'],
        ['δ(x_H,y_H)', '加热针位置的 Dirac 源项', '—'],
    ]
)

doc.add_heading('1.2  边材热物性计算', level=2)
add_para('基于木材含水率 θ_w（m³/m³）和干密度 ρ_dry（kg/m³），采用经验公式计算：')
add_para('导热系数（Thermal conductivity）：', bold=True)
add_equation('k_sw = 0.10 + 0.20 · G_b + 0.60 · θ_w    (W/m/K)')
add_para('其中 G_b = ρ_dry / ρ_w 为基本密度比（basic specific gravity），ρ_w = 1000 kg/m³。')

add_para('容积热容（Volumetric heat capacity）：', bold=True)
add_equation('ρc_sw = ρ_dry · c_dry + ρ_w · c_w · θ_w    (J/m³/K)')
add_para('其中 c_dry = 1200 J/(kg·K)，c_w = 4186 J/(kg·K)。')

add_para('热扩散系数：', bold=True)
add_equation('D_tg = k_sw / ρc_sw,    D_ax = D_tg · α')
add_para('α = D_ax/D_tg 为热扩散各向异性比，默认值 2.0（木材沿纤维方向导热更快）。')

doc.add_heading('1.3  热速度', level=2)
add_para('将液流体积速率 v_sap（cm/h）转化为热速度 v_th（m/s）：')
add_equation('v_th = (ρ_w · c_w / ρc_sw) · v_sap')
add_para('热速度表征液流对热场的对流传输强度，是有限差分方程中对流项的系数。')

doc.add_heading('1.4  加热模型', level=2)
add_para('加热探针视为轴向无限长线热源（line heat source），焦耳热功率为：')
add_equation('P = I² R    (W)')
add_para('其中电流 I = min(I_slider, V/R)，受供电电压 V 限制。标称线热源密度：')
add_equation('q_nom = P / L    (W/m)')
add_para('L 为加热探针有效长度。')

add_para('端部效应修正（End-effect correction）：', bold=True)
add_para('实际加热探针有限长，靠近两端时三维散热使有效热源强度降低。采用指数衰减模型：')
add_equation('f_end = min(1, 1 − exp(−d/r₀) − exp(−(L−d)/r₀))')
add_para('其中 d 为测温点距探针外端的径向深度，r₀ = 5 mm 为特征衰减长度。修正后的有效热源密度：')
add_equation('q = q_nom · max(0.1, f_end)')

doc.add_heading('1.5  探针布局（Nadezhdina HFD 四针法）', level=2)
add_para('按 Nadezhdina et al. (1998) 标准布局：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(2)
run = p.add_run(
    '        T1 (+Z_ax, 上游/顺流)\n'
    '        |\n'
    '        |  Z_ax\n'
    '        |\n'
    'T3 ---- H ----          (切向 Z_tg)\n'
    '        |\n'
    '        |  Z_ax\n'
    '        |\n'
    '        T2 (−Z_ax, 下游/参考)'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_para('• H：加热探针（中心），线热源')
add_para('• T1：上游轴向测温针，位于 H 上方 +Z_ax')
add_para('• T2：下游轴向测温针（参考），位于 H 下方 −Z_ax')
add_para('• T3：切向测温针，位于 H 切向 +Z_tg')

doc.add_heading('1.6  HFD 温差指标', level=2)
add_para('从四针温度计算三个关键温差：')
add_equation('ΔT_sym = T₁ − T₂    （对称温差）')
add_equation('ΔT_as = T₃ − T₂    （非对称温差）')
add_equation('ΔT_s-a = ΔT_sym − ΔT_as = T₁ − T₃')
add_para('以及比值 R = ΔT_sym / ΔT_as，作为 K-diagram 的横轴。')

# ========== 2 数值方法 ==========
doc.add_heading('2  数值方法', level=1)

doc.add_heading('2.1  空间离散', level=2)
add_para('采用均匀正方形网格，网格间距 Δx（可调，默认 0.5 mm）。网格尺寸自动计算：')
add_equation('N = ⌊70 / Δx_mm⌋')
add_para('若 N 为偶数则加 1，确保中心网格点恰好位于域中心。探针位置映射为整数网格索引。')

doc.add_heading('2.2  稳态求解（Gauss-Seidel with SOR）', level=2)
add_para('稳态场 ∂T/∂t = 0 采用逐次超松弛（Successive Over-Relaxation, SOR）Gauss-Seidel 迭代法。')
add_para('对流项采用一阶迎风格式（first-order upwind scheme）离散：')
add_equation('a_x = D_tg / Δx²,    a_y = D_ax / Δx²,    b_v = v_th / Δx')
add_para('根据流向确定迎风系数：')
add_equation('a_y,N = a_y − min(b_v, 0),    a_y,S = a_y + max(b_v, 0)')
add_equation('a_c = 2·a_x + a_y,N + a_y,S')

add_para('Gauss-Seidel 更新公式（含源项 S）：')
add_equation('T_n = [a_x·(T_{i+1,j} + T_{i−1,j}) + a_y,N·T_{i,j+1} + a_y,S·T_{i,j−1} + S] / a_c')

add_para('SOR 加速（松弛因子 ω = 1.7）：')
add_equation('T_{i,j}^(k+1) = T_{i,j}^(k) + ω · (T_n − T_{i,j}^(k))')

add_para('收敛判据为最大修正量 max|T^(k+1) − T^(k)| < 5×10⁻⁸，最大迭代次数 60000。')
add_para('边界条件：四边均为 Dirichlet 条件，T = T_ambient。', bold=True)
add_para('暖启动（Warm start）：在日变化回放中，前一时刻的稳态解作为下一时刻的初始猜测，大幅加速收敛。', bold=True)

doc.add_heading('2.3  瞬态求解（显式 FTCS + Upwind）', level=2)
add_para('瞬态模式采用显式前向时间-中心空间（Forward Time, Central Space, FTCS）格式，对流项同样采用一阶迎风：')
add_equation('T_{i,j}^{n+1} = T_{i,j}^n + r_x·(T_{i+1,j} − 2T_{i,j} + T_{i−1,j}) + r_y·(T_{i,j+1} − 2T_{i,j} + T_{i,j−1}) − C_v·(upwind diff) + S·Δt')
add_para('其中：')
add_equation('r_x = D_tg·Δt / Δx²,    r_y = D_ax·Δt / Δx²,    C_v = v_th·Δt / Δx')
add_para('迎风项根据 C_v 符号选择：C_v ≥ 0 时取 T_{i,j} − T_{i,j−1}；C_v < 0 时取 T_{i,j+1} − T_{i,j}。')

add_para('时间步长由 CFL 稳定性条件自动确定：', bold=True)
add_equation('Δt = min( 0.2·Δx²/max(D_tg,D_ax),  0.4·Δx/|v_th|,  0.5s )')

doc.add_heading('2.4  数值格式选择理由', level=2)
add_table(
    ['方面', '选择', '理由'],
    [
        ['对流离散', '一阶迎风', '无条件稳定（无非物理振荡），对 HFD 温差场足够精确'],
        ['稳态求解', 'Gauss-Seidel + SOR', '内存占用极小（原地更新），ω=1.7 接近最优，收敛快'],
        ['瞬态求解', '显式 FTCS', '实现简单，配合 CFL 自适应步长保证稳定性'],
        ['数据结构', 'Float64Array 一维展平', '高性能连续内存访问，适合浏览器 JIT 优化'],
    ]
)

# ========== 3 数据处理 ==========
doc.add_heading('3  数据处理与分析', level=1)

doc.add_heading('3.1  K 值计算（Nadezhdina 方法）', level=2)
add_para('K 值是 HFD 方法的核心校正参数，代表零液流条件下的自然温度梯度贡献。')
add_para('算法流程：', bold=True)
add_para('1. 筛选低流量点：按 |R| = |ΔT_sym/ΔT_as| 从小到大排序，取前 15%（最少 3 个点）')
add_para('2. 线性回归：以 R 为自变量，ΔT_as 为因变量，求最小二乘回归')
add_para('3. 提取截距：K = 回归截距（R=0 时的 ΔT_as 值）')
add_para('回归公式（普通最小二乘法）：')
add_equation('K = (Σyᵢ·Σxᵢ² − Σxᵢ·Σxᵢyᵢ) / (n·Σxᵢ² − (Σxᵢ)²)')
add_para('其中 xᵢ = Rᵢ，yᵢ = ΔT_as,i。当分母趋近于零时，退化为均值估计 K = ȳ。')
add_para('4. 修正温差：K + ΔT_s-a 用于 K-diagram 中校正后的温差线')

doc.add_heading('3.2  日变化液流曲线生成', level=2)
add_para('采用升余弦平方（raised-cosine-squared）模型生成平滑的日液流变化曲线：')
add_equation('f(h) = cos²(π·(h − h_peak) / (2W))    当 |h − h_peak| < W，否则为 0')
add_equation('v_sap(h) = v_min + (v_peak − v_min) · f(h)')
add_para('其中 W = 7.5 h 为半宽度，h_peak 为峰值时刻。时间分辨率为 10 min/点，一天 144 点。')

add_para('天气修正模式：', bold=True)
add_table(
    ['模式', '修正方法'],
    [
        ['Sunny', '无修正，标准 cos² 钟形曲线'],
        ['Cloudy', '叠加高斯凹陷：v ← v − 0.3·(v_peak−v_min)·exp[−(h−12.5)²/(2×1.2²)]·f(h)，模拟正午云遮'],
        ['Drought', 'h > 11h 后指数衰减：f ← f·exp[−0.25·(h−11)]，模拟干旱提前关闭气孔'],
    ]
)
add_para('噪声模型：加性均匀随机噪声 ε = v_peak · σ_noise · U(−1,1)。')

doc.add_heading('3.3  外部数据上传', level=2)
add_para('支持 CSV/TXT 格式上传实测液流数据：')
add_para('• 解析格式：逗号/制表符/分号分隔，第一列为时间戳，第二列为 v_sap（cm/h）')
add_para('• 时间戳支持：完整日期时间（2024-01-15 08:30）或仅时间（08:30），自动处理午夜跨越')
add_para('• 数据约束：最少 144 行（1 天），最多 14400 行')
add_para('• 多日支持：自动检测多天数据，时间轴显示天数标签')

# ========== 4 K-Diagram ==========
doc.add_heading('4  K-Diagram 可视化', level=1)

doc.add_heading('4.1  坐标系', level=2)
add_para('K-diagram（Nadezhdina 2018）以温差比值为横轴，各温差为纵轴：')
add_para('• x 轴：R = ΔT_sym / ΔT_as')
add_para('• y 轴：温度（°C）')

doc.add_heading('4.2  绘制内容', level=2)
add_table(
    ['图层', '数据', '颜色', '含义'],
    [
        ['散点 ΔT_sym', 'R vs T₁−T₂', '黑色', '对称温差随流量变化'],
        ['散点 ΔT_as', 'R vs T₃−T₂', '粉色', '非对称温差'],
        ['散点 ΔT_s-a', 'R vs T₁−T₃', '灰色', '对称减非对称'],
        ['散点 K+ΔT_s-a', 'R vs K+(T₁−T₃)', '绿色', 'K 校正后温差'],
        ['K 水平线', 'y = K', '红色虚线', '零流量截距'],
        ['零线', 'y = 0', '蓝色虚线', '温度零参考'],
    ]
)

doc.add_heading('4.3  物理意义', level=2)
add_para('• 零流量时（R ≈ 0）：ΔT_sym ≈ 0，ΔT_as 和 K+ΔT_s-a 应在 K 值处交汇')
add_para('• 高流量时（R ≫ 0）：ΔT_sym 增大，热场明显向流向偏移')
add_para('• K+ΔT_s-a 与 ΔT_as 的交点验证 K 值估计的准确性')

# ========== 5 热场可视化 ==========
doc.add_heading('5  热场可视化', level=1)

doc.add_heading('5.1  色图（Colormap）', level=2)
add_para('采用改进的 Jet 色图，将温升 ΔT = T − T_ambient 归一化到 [0, 1] 后映射：')
add_table(
    ['范围 t', '颜色过渡'],
    [
        ['[0, 0.05)', '暗底 → 深蓝（从黑色渐入，避免零值处突变）'],
        ['[0, 0.125)', '深蓝 → 蓝'],
        ['[0.125, 0.375)', '蓝 → 青'],
        ['[0.375, 0.625)', '青 → 黄'],
        ['[0.625, 0.875)', '黄 → 红'],
        ['[0.875, 1.0]', '红 → 深红'],
    ]
)

doc.add_heading('5.2  双线性插值渲染', level=2)
add_para('为避免网格像素化伪影，渲染时对每个屏幕像素进行双线性插值（bilinear interpolation）：')
add_equation('T(x,y) = (1−f_x)(1−f_y)·T₀₀ + f_x(1−f_y)·T₁₀ + (1−f_x)f_y·T₀₁ + f_x·f_y·T₁₁')
add_para('通过 ImageData 逐像素写入 RGBA 数据，利用 HTML5 Canvas 2D API 的 putImageData() 高效渲染。')

doc.add_heading('5.3  等温线（Isotherms）', level=2)
add_para('在色图上叠加等温线，共 12 条，均匀分布于 [T_ambient, T_ambient + ΔT_max]。')
add_para('等温线提取采用 Marching Squares 算法：逐网格单元检查四角温度值与等值面的交叉关系，通过 4-bit 索引（16 种情况）确定等温线段的起止端点，线性插值确定精确位置。')

doc.add_heading('5.4  Canvas 自适应尺寸', level=2)
add_para('热场画布尺寸根据浏览器窗口自适应计算：')
add_equation('sz = max(180, min(W_available, H_available, 560))')
add_para('窗口大小变化时自动调整并重绘。')

# ========== 6 时间序列图 ==========
doc.add_heading('6  时间序列图', level=1)

doc.add_heading('6.1  双 y 轴设计', level=2)
add_para('• 左 y 轴：温差（°C），绘制 ΔT_sym、ΔT_as、ΔT_s-a 三条曲线')
add_para('• 右 y 轴：液流速率 v_sap（cm/h），绘制为绿色填充面积图')
add_para('• x 轴：时间（h），自动适应单日或多日数据')

doc.add_heading('6.2  刻度自适应', level=2)
add_para('坐标刻度间隔通过对数量级自适应算法确定：')
add_equation('step = 0.2p (当 r/p < 2),  0.5p (当 r/p < 5),  p (其他)')
add_equation('其中 p = 10^⌊log₁₀(range)⌋')

# ========== 7 技术架构 ==========
doc.add_heading('7  技术架构', level=1)

doc.add_heading('7.1  实现环境', level=2)
add_table(
    ['项目', '技术'],
    [
        ['语言', 'JavaScript (ES6+)'],
        ['渲染', 'HTML5 Canvas 2D API'],
        ['数值数组', 'Float64Array（IEEE 754 双精度）'],
        ['动画', 'requestAnimationFrame（瞬态）/ setTimeout（日变化回放）'],
        ['部署', '单页面 HTML，GitHub Pages 静态托管'],
        ['依赖', '零外部依赖，完全自包含'],
    ]
)

doc.add_heading('7.2  性能优化', level=2)
add_para('• 一维数组展平：二维网格以行优先方式存储在一维 Float64Array 中，索引 idx = i × N_Y + j，利用连续内存局部性加速缓存命中')
add_para('• SOR 暖启动：日变化模式中复用前一稳态解作为初始猜测，相邻时刻温度场变化小，收敛所需迭代数大幅减少')
add_para('• 批处理子步：瞬态模式每帧计算多个物理时间步后再渲染一次，平衡物理精度与帧率')
add_para('• 降采样绘制：日变化预览曲线在数据点多于画布像素时自动降采样')

doc.add_heading('7.3  输入参数范围', level=2)
add_table(
    ['参数', '最小值', '最大值', '默认值', '单位'],
    [
        ['Z_ax', '5', '30', '15', 'mm'],
        ['Z_tg', '2', '15', '5', 'mm'],
        ['Heater L', '20', '120', '100', 'mm'],
        ['Supply V', '1.0', '12.0', '6.0', 'V'],
        ['Current I', '20', '150', '50', 'mA'],
        ['Resistance R', '40', '300', '120', 'Ω'],
        ['θ_w', '0.10', '0.80', '0.40', 'm³/m³'],
        ['ρ_dry', '200', '900', '500', 'kg/m³'],
        ['D_ax/D_tg', '1.0', '4.0', '2.0', '—'],
        ['v_sap', '−20', '150', '0', 'cm/h'],
        ['T_ambient', '0', '40', '20', '°C'],
        ['Grid Δ', '0.3', '1.0', '0.5', 'mm'],
    ]
)

# ========== 8 参考文献 ==========
doc.add_heading('8  参考文献', level=1)
refs = [
    'Nadezhdina, N., Čermák, J., & Nadezhdin, V. (1998). Heat field deformation method for sap flow measurements. Proceedings of the 4th International Workshop on Measuring Sap Flow in Intact Plants, 72–92.',
    'Nadezhdina, N., Vandegehuchte, M. W., & Steppe, K. (2012). Sap flux density measurements based on the heat field deformation method. Trees, 26, 1439–1448.',
    'Nadezhdina, N. (2018). Revisiting the Heat Field Deformation (HFD) method for measuring sap flow. iForest – Biogeosciences and Forestry, 11, 118–130.',
    'ICT International. HFD8-100 Heat Field Deformation Sap Flux Meter. Product specifications.',
    'Patankar, S. V. (1980). Numerical Heat Transfer and Fluid Flow. Hemisphere Publishing.',
    'LeVeque, R. J. (2007). Finite Difference Methods for Ordinary and Partial Differential Equations. SIAM.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}]  {ref}')
    run.font.size = Pt(10)

# Save
out = '/home/user/HFD-simulator/METHODS.docx'
doc.save(out)
print(f'Saved to {out}')
